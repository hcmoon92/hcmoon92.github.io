import sys
from docx import Document
import re

def extract_major_topics(doc_path, output_path):
    try:
        # Load the .doc file
        doc = Document(doc_path)
        
        # Open the output file in write mode
        with open(output_path, 'w', encoding='utf-8') as md_file:
            md_file.write("# Extracted Major Topics\n\n")
            
            for paragraph in doc.paragraphs:
                # Consider paragraphs with specific styles as major topics
                if paragraph.style.name in ["Heading 1", "Title", "Subtitle"]:
                    md_file.write(f"## {paragraph.text.strip()}\n\n")
        
        print(f"Major topics successfully extracted to {output_path}")
    except Exception as e:
        print(f"An error occurred: {e}")


def heading_to_markdown(heading_text, level):
    """Converts heading text to markdown format based on heading level."""
    prefix = "#" * level  # Markdown heading level
    return f"{prefix} {heading_text}"

def process_docx_to_markdown(input_path, output_path):
    """Reads a .docx file and generates a structured markdown file."""
    try:
        document = Document(input_path)
        markdown_lines = []

        # Loop through paragraphs in the document
        for paragraph in document.paragraphs:
            if paragraph.style.name.startswith("Heading"):
                # Extract heading level (e.g., Heading 1 -> 1, Heading 2 -> 2)
                try:
                    level = int(paragraph.style.name.split(" ")[-1])
                except ValueError:
                    level = 1  # Default to level 1 for undefined heading levels

                # Convert to markdown format
                markdown_lines.append(heading_to_markdown(paragraph.text.strip(), level))
            elif paragraph.style.name in ["Title", "Subtitle"]:
                # Handle special styles like Title or Subtitle
                markdown_lines.append(f"## {paragraph.text.strip()}")  # Subtitle -> Markdown H2
            else:
                # Regular text or unstyled content
                markdown_lines.append(paragraph.text.strip())

        # Write to markdown file
        with open(output_path, "w", encoding="utf-8") as markdown_file:
            markdown_file.write("\n\n".join(markdown_lines))

        print(f"Successfully generated markdown: {output_path}")

    except Exception as e:
        print(f"An error occurred: {e}")


# 5 and 5.1
def find_toc_title(doc_path):
    doc = Document(doc_path)
    toc_found = False
    toc = []

    # Regular expression to match lines with numbers like x, x.xx, or x.x or alphanumeric like D.1
    toc_pattern = re.compile(r'^[A-Za-z]?\d+(\.\d{1,2})?\s')

    for paragraph in doc.paragraphs:
        # Assuming TOC starts with "Contents" or is styled as a specific Heading
        if "Contents" in paragraph.text or paragraph.style.name.startswith("Heading"):
            toc_found = True

        if toc_found:
            # Collect TOC content until it reaches the end of TOC section
            if paragraph.text.strip() == "":
                break  # End of TOC section, assuming blank line marks the end
            # Only append lines that match the number pattern (x, x.xx, x.x, or alphanumeric like D.1)
            if toc_pattern.match(paragraph.text.strip()):
                toc.append(paragraph.text.strip())

    return toc

# 5 and 5.1
def process_docx_to_markdown_toc(input_path, output_path):
    doc = Document(input_path)
    toc_found = False
    toc = []

    # Regular expression to match lines with numbers like x, x.xx, or x.x or alphanumeric like D.1
    toc_pattern = re.compile(r'^[A-Za-z]?\d+(\.\d{1,2})?\s')

    for paragraph in doc.paragraphs:
        # Assuming TOC starts with "Contents" or is styled as a specific Heading
        if "Contents" in paragraph.text or paragraph.style.name.startswith("Heading"):
            toc_found = True

        if toc_found:
            # Collect TOC content until it reaches the end of TOC section
            if paragraph.text.strip() == "":
                break  # End of TOC section, assuming blank line marks the end
            # Only append lines that match the number pattern (x, x.xx, x.x, or alphanumeric like D.1)
            if toc_pattern.match(paragraph.text.strip()):
                toc.append(paragraph.text.strip())
                
    # # Write to markdown file
    # with open(output_path, "w", encoding="utf-8") as markdown_file:
    #     markdown_file.write("\n".join(toc))

    # print(f"Successfully generated markdown: {output_path}")
    
    return toc

def convert_to_markdown(toc_list):
    markdown_toc = "# Table of Contents\n\n"
    
    current_section = None  # Track the current main section
    for line in toc_list:
        # Match for main sections (e.g., 1, 2, 3, etc.)
        main_section_match = re.match(r"(\d+)(?:\s+)([A-Za-z0-9\s\-]+)(\s+\d+)", line)
        # Match for subsections (e.g., 3.1, 5.1, etc.)
        subsection_match = re.match(r"(\d+\.\d+)(?:\s+)([A-Za-z0-9\s\-]+)(\s+\d+)", line)
        
        if main_section_match:
            # Format the main section
            section_number = main_section_match.group(1)
            section_title = main_section_match.group(2).strip()
            page_number = main_section_match.group(3).strip()
            markdown_toc += f"1. **{section_number}. {section_title}** {' ' * (60 - len(section_number + section_title))}{page_number}\n"
            current_section = section_number  # Update the current main section
        elif subsection_match:
            # Format the subsection under the parent section
            subsection_number = subsection_match.group(1)
            subsection_title = subsection_match.group(2).strip()
            page_number = subsection_match.group(3).strip()
            
            # Check if the subsection belongs to the current main section
            if subsection_number.startswith(current_section + "."):
                # Add indentation for subsections
                markdown_toc += f"  - {subsection_number}. **{subsection_title}** {' ' * (60 - len(subsection_number + subsection_title))}{page_number}\n"
    
    return markdown_toc


def extract_title_from_first_page(doc_path):
    # Load the document
    doc = Document(doc_path)
    
    # Extract text from the first page (or top paragraphs)
    first_page_content = []
    for paragraph in doc.paragraphs:
        first_page_content.append(paragraph.text)
        
        # Stop after reading the first page (usually the first few paragraphs)
        # You can adjust this based on your document structure
        if len(first_page_content) > 5:  # Adjust based on expected number of paragraphs per page
            break

    #print(first_page_content)
    return first_page_content


def get_header(doc_path):
    doc = Document(doc_path)
    headers = doc.sections[0].header
    header_content = [para.text for para in headers.paragraphs]
    return header_content

def main():
    if len(sys.argv) != 2:
        print("Usage: python reader.py <input.doc>")
        return
    
    input_path = sys.argv[1]
    #output_path = sys.argv[2]
    
    #extract_major_topics(input_path, output_path)
    #process_docx_to_markdown(input_path, output_path)

    #toc_content = find_toc_title(input_path)
    # Print TOC content
    #print("\n".join(toc_content))
    title = extract_title_from_first_page(input_path)
    # Find the index of the first occurrence of '\t'
    #tab_index = title.index("\t")
    # Slice the list up to the first '\t'
    #print(title)
    
    sliced_title = title[:4]
    sliced_title.append('\n')
    
    spec_number = title[0]
    spec_name = title[4]

    #header = get_header(input_path)
    #print(header)
    
    if spec_number :
        # Generate the output filename using spec_number and spec_name
        output_filename = f"{spec_number}-{spec_name}.md"
        output_path = output_filename.replace(" ", "_")  # Optional: replace spaces with underscores
    else : 
        # Generate the output filename using spec_number and spec_name
        output_filename = f"spec.md"
        output_path = output_filename.replace(" ", "_")  # Optional: replace spaces with underscores
    

    toc = process_docx_to_markdown_toc(input_path, output_path)
    markdown_output = convert_to_markdown(toc)
    print(markdown_output)

    # Write to markdown file
    with open(output_path, "w", encoding="utf-8") as markdown_file:
        markdown_file.writelines('\n'.join(sliced_title))  # Join list elements with newlines and write them
        markdown_file.write(markdown_output)

    print(f"Successfully generated markdown: {output_path}")
    
if __name__ == "__main__":
    main()
